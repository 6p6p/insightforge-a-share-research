"""DOCX renderer (stage 6C spec O)：python-docx，纯函数，不查 DB。

结构：标题 + 元数据 + section（Heading1）+ 段落（正文 + [n] 标记）+ 证据附录
（E1..En）+ audit_note。段落 `text` 原样（**不改写句子**），编号标记由
citation_numbers 追加。输出为内存 bytes（reopen-able）。

**字节确定性（v2）**：
1. core properties：显式写死 created/modified（固定值，与真实创建时间无关）；
2. **OOXML ZIP normalize**：python-docx 保存时 `ZipInfo.date_time` 取保存时刻
   （秒级），同一 pack 在不同 wall-clock 时间渲染会得到不同字节（sha256 不等）
   ——只固定 core properties 不够。`_normalize_docx_zip` 把容器重写为规范形式：
   entry 按 filename 稳定排序 + 每个 ZipInfo 固定 date_time / compress_type /
   create_system / external_attr / extra / comment（等元数据），**不改 entry
   content 字节**。输出仍是合法 DOCX（python-docx 重开只依赖 XML part 内容）。
"""

from datetime import datetime
from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo
from zoneinfo import ZoneInfo

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

from app.report_export.contracts import ExportReportPack

# 固定 core properties（字节确定性；与真实创建时间无关）。
_FIXED_CREATED = datetime(1970, 1, 1, tzinfo=ZoneInfo("UTC"))

# 合法 ZIP（DOS）时间戳最小值：所有 entry date_time 固定为该值（字节确定性，
# 不依赖系统时区 / 保存时刻）。
_ZIP_FIXED_TIMESTAMP = (1980, 1, 1, 0, 0, 0)
# entry 固定 unix mode（regular file 0600）。create_system=3（Unix）时外部属性
# 高 16 位为 mode 位——统一固定，消除跨系统 DOS/Unix 与 mtime 位漂移。
_ZIP_FIXED_EXTERNAL_ATTR = 0o100600 << 16

# 中文正文西文字体 + East Asian 字体（Word 渲染时按名称解析，text 提取不受影响）。
_LATIN_FONT = "Calibri"
_EAST_ASIAN_FONT = "宋体"

_DATE_FMT = "%Y-%m-%d"


def render_docx(pack: ExportReportPack) -> bytes:
    document = Document()

    cp = document.core_properties
    cp.title = f"{pack.company_name} 基本面研究报告"
    cp.subject = "证据驱动基本面研究（InsightForge）"
    cp.author = "InsightForge"
    cp.created = _FIXED_CREATED
    cp.modified = _FIXED_CREATED

    _set_default_east_asian(document)

    security = f"（{pack.security_code}）" if pack.security_code else ""
    document.add_heading(f"{pack.company_name} 基本面研究报告{security}", level=0)
    _add_meta(document, "研究问题", pack.research_question or "—")
    _add_meta(document, "研究区间", _research_window(pack))
    _add_meta(document, "分析基准日", pack.analysis_as_of.isoformat())

    for section in pack.sections:
        document.add_heading(section.title, level=1)
        for paragraph in section.paragraphs:
            markers = "".join(f"[{n}]" for n in paragraph.citation_numbers)
            _add_body(document, f"{paragraph.text}{markers}")

    if pack.citations:
        document.add_heading("证据附录", level=1)
        for citation in pack.citations:
            document.add_heading(f"E{citation.number} ｜ {citation.provider_label or '—'}", level=2)
            _add_citation_lines(document, citation)

    if pack.audit_note:
        paragraph = document.add_paragraph(pack.audit_note)
        paragraph.style = document.styles["Quote"]

    buffer = BytesIO()
    document.save(buffer)
    return _normalize_docx_zip(buffer.getvalue())


def _normalize_docx_zip(data: bytes) -> bytes:
    """确定性 OOXML ZIP normalize（spec A2；renderer v2）。

    python-docx 保存时 `ZipInfo.date_time` 取保存时刻（秒级），同一 pack 在
    不同 wall-clock 时间渲染会得到不同字节——只固定 core properties 不够。
    这里把 OOXML 容器重写为规范形式：

    - entry 按 filename 稳定排序（容器内顺序与 python-docx 内部次序解耦）；
    - 每个 ZipInfo 固定 date_time / compress_type / create_system /
      external_attr / extra / comment（及 create/extract version、
      internal_attr）——任何会导致 byte drift 的 metadata 全部钉死；
    - **不改变 entry content 字节**：同内容 deflate 在同一 zlib 实现下字节
      确定（`verify_export_integrity` 按 content_sha256 比对归档字节）。

    输出仍是合法 DOCX：python-docx 打开只依赖 [Content_Types].xml 与
    word/document.xml 的 content，不受 metadata / entry 顺序影响。
    """
    output = BytesIO()
    with ZipFile(BytesIO(data), "r") as zipped:
        infos = sorted(zipped.infolist(), key=lambda item: item.filename)
        with ZipFile(output, "w", compression=ZIP_DEFLATED) as out:
            for info in infos:
                content = zipped.read(info.filename)
                fixed = ZipInfo(info.filename)
                fixed.date_time = _ZIP_FIXED_TIMESTAMP
                fixed.compress_type = ZIP_DEFLATED
                fixed.comment = b""
                fixed.extra = b""
                fixed.create_system = 3
                fixed.create_version = 20
                fixed.extract_version = 20
                fixed.internal_attr = 0
                fixed.external_attr = _ZIP_FIXED_EXTERNAL_ATTR
                out.writestr(fixed, content)
    return output.getvalue()


def _set_default_east_asian(document: Document) -> None:
    """Normal 样式同时设置西文 + East Asian 字体（中文正文可读）。"""
    normal = document.styles["Normal"]
    normal.font.name = _LATIN_FONT
    normal.font.size = Pt(11)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), _EAST_ASIAN_FONT)


def _research_window(pack: ExportReportPack) -> str:
    """研究区间（YYYY-MM-DD ~ YYYY-MM-DD；缺失 → "—"）。"""
    if pack.research_start_date is None or pack.research_end_date is None:
        return "—"
    return (
        f"{pack.research_start_date.isoformat()} ~ {pack.research_end_date.isoformat()}"
    )


def _add_meta(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"{label}：")
    run.bold = True
    paragraph.add_run(value)


def _add_body(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.paragraph_format.first_line_indent = Pt(22)


def _add_citation_lines(document: Document, citation) -> None:
    _add_bullet(document, "证据陈述", citation.statement or "—")
    if citation.quote_text:
        _add_bullet(document, "引用原文", f"「{citation.quote_text}」")
    # clean projection：只保留人类可读来源引用（不输出 provider code）。
    _add_bullet(document, "提供方", citation.provider_label or "—")
    _add_bullet(document, "来源", citation.title or "—")
    if citation.source_url:
        _add_bullet(document, "原始网页", citation.source_url)
    if citation.published_at is not None:
        _add_bullet(document, "发布", citation.published_at.strftime(_DATE_FMT))
    if citation.fetched_at is not None:
        _add_bullet(document, "获取", citation.fetched_at.strftime(_DATE_FMT))
    if citation.page_number is not None:
        _add_bullet(document, "定位", f"第 {citation.page_number} 页")
    if (
        citation.indicator is not None
        or citation.geography is not None
        or citation.period is not None
    ):
        parts = []
        if citation.indicator:
            parts.append(f"指标 {citation.indicator}")
        if citation.geography:
            parts.append(f"地域 {citation.geography}")
        if citation.period:
            parts.append(f"观测期 {citation.period}")
        _add_bullet(document, "宏观", " ｜ ".join(parts))


def _add_bullet(document: Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"{label}：")
    run.bold = True
    paragraph.add_run(value)
    paragraph.paragraph_format.left_indent = Pt(12)
