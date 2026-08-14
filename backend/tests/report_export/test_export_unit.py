"""Pure unit tests for deterministic report export (stage 6C spec I/J/L/M/N/O).

无 DB / 无 LLM / 无网络——只测确定性代码本身：
- `build_export_report_pack`：引用编号（spec J，section_order → paragraph_index
  → evidence_card_ids 首次出现 → E1..En，同一卡恒同号）+ company fallback +
  audit_note 透传；
- 三个 renderer：字节确定性（同 pack → 同 bytes）+ Markdown 结构 + DOCX
  reopen-able + PDF（%PDF magic + pdfplumber 中文可读）；
- `compute_export_input_fingerprint`：稳定（同输入 → 同指纹）+ 敏感（任一
  format / report / renderer / audit / pack 变化 → 新指纹）；
- `ExportArtifactStore`：内容寻址去重 + 空字节拒绝 + 路径守卫（无 caller-
  controlled path 逃逸）。
"""

import hashlib
import time
import zipfile
from datetime import date, datetime
from io import BytesIO
from uuid import UUID, uuid4

import pytest
from docx import Document

from app.db.models.research_task import ResearchTaskModel
from app.report.contracts import VerifiedReport
from app.report_export.contracts import (
    RENDERER_NAME_BY_FORMAT,
    RENDERER_VERSION_BY_FORMAT,
    ExportCitation,
    ExportParagraph,
    ExportReportPack,
    ExportSection,
    compute_export_input_fingerprint,
)
from app.report_export.errors import ExportArtifactNotFound
from app.report_export.pack import (
    AUDIT_NOTE_HUMAN_APPROVED,
    ExportCardDetail,
    build_export_report_pack,
)
from app.report_export.renderers import render_export
from app.report_export.renderers.docx import render_docx
from app.report_export.renderers.markdown import render_markdown
from app.report_export.renderers.pdf import render_pdf
from app.schemas.citation import CitationLocator, DocumentProvenance
from app.storage.export_store import ExportArtifactStore, InvalidStorageKey

# ---------------------------------------------------------------- fixtures


def _doc_provenance(*, page_number: int | None = None, xpath: str | None = None):
    return DocumentProvenance(
        origin_type="document_chunk",
        source_id=uuid4(),
        provider_key="sse",
        provider_label="上交所",
        title="2024年年度报告",
        source_url="https://www.sse.com.cn/a.pdf",
        published_at=datetime(2025, 3, 1),
        authority_tier=1,
        document_type="annual_report",
        raw_artifact_id=uuid4(),
        media_type="application/pdf",
        parsed_source_id=uuid4(),
        chunk_id=uuid4(),
        locator=CitationLocator(locator_type="pdf_page", page_number=page_number)
        if page_number
        else CitationLocator(locator_type="html_dom", xpath=xpath),
        context_text="上下文纯文本",
        quote_text="2024年公司经营现金流净额同比增长20%。",
    )


def _card_detail(card_id, statement: str) -> ExportCardDetail:
    return ExportCardDetail(
        evidence_card_id=card_id,
        statement=statement,
        quote_text=None,
        origin_type="document_chunk",
    )


def _sample_pack(*, audit_note: str | None = None) -> ExportReportPack:
    return ExportReportPack(
        export_schema_version=1,
        task_id=uuid4(),
        report_id=uuid4(),
        analysis_as_of=date(2024, 1, 1),
        company_name="贵州茅台",
        security_code="600519",
        research_question="600519 基本面研究",
        sections=(
            ExportSection(
                section_id="s1",
                title="业务概况",
                paragraphs=(
                    ExportParagraph(
                        text="公司经营现金流净额同比增长20%。", citation_numbers=(1, 2)
                    ),
                    ExportParagraph(text="第二段无引用。", citation_numbers=()),
                ),
            ),
        ),
        citations=(
            ExportCitation(
                number=1,
                evidence_card_id=uuid4(),
                statement="经营现金流同比增长20%",
                quote_text="原文A",
                origin_type="document_chunk",
                provider_key="sse",
                provider_label="上交所",
                title="2024年年度报告",
                source_url="https://www.sse.com.cn/a.pdf",
                published_at=datetime(2025, 3, 1),
                page_number=12,
            ),
            ExportCitation(
                number=2,
                evidence_card_id=uuid4(),
                statement="宏观GDP增速回升",
                origin_type="macro_observation",
                provider_key="nbs",
                provider_label="国家统计局",
                title="GDP同比",
                indicator="GDP同比",
                geography="全国",
                period="2024Q1",
                fetched_at=datetime(2024, 2, 1),
            ),
        ),
        audit_note=audit_note,
        report_fingerprint="f" * 64,
    )


def _fingerprint_args(
    pack: ExportReportPack,
    *,
    format: str = "markdown",
    renderer_version: int | None = None,
    **overrides,
) -> dict:
    args = {
        "export_schema_version": pack.export_schema_version,
        "task_id": pack.task_id,
        "report_id": pack.report_id,
        "report_fingerprint": pack.report_fingerprint,
        "check_result_id": UUID("00000000-0000-0000-0000-000000000001"),
        "check_fingerprint": "c" * 64,
        "audit_id": UUID("00000000-0000-0000-0000-000000000002"),
        "audit_fingerprint": "a" * 64,
        "human_decision_id": None,
        "decision_fingerprint": None,
        "format": format,
        "renderer_name": RENDERER_NAME_BY_FORMAT[format],
        "renderer_version": RENDERER_VERSION_BY_FORMAT[format]
        if renderer_version is None
        else renderer_version,
        "pack_identity": pack.to_identity_dict(),
    }
    args.update(overrides)
    return args


# ZIP 时间戳固定值：合法 DOS 时间最小值（1980-01-01 00:00:00），与 renderer 内
# `_ZIP_FIXED_TIMESTAMP` 同步（ZIP normalize 后所有 entry date_time == 该值）。
_ZIP_FIXED_TIMESTAMP = (1980, 1, 1, 0, 0, 0)


# ---------------------------------------------------------------- pack numbering (spec J)


def test_pack_numbering_first_occurrence() -> None:
    """section_order → paragraph_index → evidence_card_ids 首次出现 → E1..En。"""
    card_a, card_b, card_c = uuid4(), uuid4(), uuid4()
    payload = {
        "sections": [
            {
                "section_id": "s1",
                "title": "业务概况",
                "paragraphs": [
                    {"text": "para1", "evidence_card_ids": [str(card_b)]},
                    {"text": "para2", "evidence_card_ids": [str(card_a), str(card_b)]},
                ],
            },
            {
                "section_id": "s2",
                "title": "风险",
                "paragraphs": [
                    {"text": "para3", "evidence_card_ids": [str(card_a), str(card_c)]},
                    {"text": "para4", "evidence_card_ids": []},
                ],
            },
        ]
    }
    report = VerifiedReport(
        report_id=uuid4(),
        outline_id=uuid4(),
        company_id=uuid4(),
        research_question_sha256="x",
        analysis_as_of=date(2024, 1, 1),
        report_schema_version=1,
        report_fingerprint="f" * 64,
        report_payload=payload,
        verified_outline=object(),  # pack builder 不读上游投影
        verified_drafts=(),
    )
    task = ResearchTaskModel(
        task_id=uuid4(),
        company_query="600519",
        research_start_date=date(2023, 1, 1),
        research_end_date=date(2026, 12, 31),
        modules=["company_profile"],
        questions=["600519 基本面研究"],
        require_plan_approval=False,
    )
    cards_by_id = {
        card_a: _card_detail(card_a, "A statement"),
        card_b: _card_detail(card_b, "B statement"),
        card_c: _card_detail(card_c, "C statement"),
    }
    provenance = {cid: _doc_provenance() for cid in (card_a, card_b, card_c)}

    pack = build_export_report_pack(
        verified_report=report,
        task=task,
        company=None,
        cards_by_id=cards_by_id,
        provenance_by_card=provenance,
        audit_note=None,
    )

    # 首次出现顺序：card_b → E1，card_a → E2，card_c → E3。
    assert pack.sections[0].paragraphs[0].citation_numbers == (1,)  # card_b
    assert pack.sections[0].paragraphs[1].citation_numbers == (2, 1)  # card_a, card_b
    assert pack.sections[1].paragraphs[0].citation_numbers == (2, 3)  # card_a, card_c
    assert pack.sections[1].paragraphs[1].citation_numbers == ()  # 无引用

    # 附录按编号升序；同一 evidence 恒同号。
    assert [c.number for c in pack.citations] == [1, 2, 3]
    assert [c.evidence_card_id for c in pack.citations] == [card_b, card_a, card_c]

    # 元数据：company fallback + 原样字段。
    assert pack.company_name == "600519"  # company=None → task.company_query
    assert pack.security_code is None
    assert pack.report_id == report.report_id
    assert pack.analysis_as_of == report.analysis_as_of
    # 正文 text 原样（不改写句子）。
    assert pack.sections[0].paragraphs[0].text == "para1"


def test_pack_audit_note_and_company() -> None:
    """company 存在时用 short_name；audit_note（人工批准路径）透传。"""
    card_a = uuid4()
    payload = {
        "sections": [
            {
                "section_id": "s1",
                "title": "业务",
                "paragraphs": [{"text": "p", "evidence_card_ids": [str(card_a)]}],
            }
        ]
    }
    report = VerifiedReport(
        report_id=uuid4(),
        outline_id=uuid4(),
        company_id=uuid4(),
        research_question_sha256="x",
        analysis_as_of=date(2024, 1, 1),
        report_schema_version=1,
        report_fingerprint="f" * 64,
        report_payload=payload,
        verified_outline=object(),
        verified_drafts=(),
    )
    task = ResearchTaskModel(
        task_id=uuid4(),
        company_query="600519",
        research_start_date=date(2023, 1, 1),
        research_end_date=date(2026, 12, 31),
        modules=["company_profile"],
        questions=[],
        require_plan_approval=False,
    )
    cards_by_id = {card_a: _card_detail(card_a, "A")}
    provenance = {card_a: _doc_provenance()}

    pack = build_export_report_pack(
        verified_report=report,
        task=task,
        company=None,
        cards_by_id=cards_by_id,
        provenance_by_card=provenance,
        audit_note=AUDIT_NOTE_HUMAN_APPROVED,
    )
    assert pack.audit_note == AUDIT_NOTE_HUMAN_APPROVED
    # questions 空 → 研究问题回退到 company_query。
    assert pack.research_question == "600519"


# ---------------------------------------------------------------- renderers (spec O)


def test_render_markdown_deterministic_and_structure() -> None:
    pack = _sample_pack(audit_note=AUDIT_NOTE_HUMAN_APPROVED)
    first = render_markdown(pack)
    assert render_markdown(pack) == first, "同 pack → 同 bytes（字节确定性）"

    text = first.decode("utf-8")
    assert "# 贵州茅台 基本面研究报告（600519）" in text
    assert "- 研究问题：600519 基本面研究" in text
    # 正文 + [n] 标记原样。
    assert "公司经营现金流净额同比增长20%。[1][2]" in text
    assert "第二段无引用。" in text
    # 附录 E1..En。
    assert "### E1 ｜ 上交所" in text
    assert "### E2 ｜ 国家统计局" in text
    assert "- 原始网页：https://www.sse.com.cn/a.pdf" in text
    assert "- 定位：第 12 页" in text
    assert "- 宏观：指标 GDP同比 ｜ 地域 全国 ｜ 观测期 2024Q1" in text
    # audit_note（人工批准路径）引用块。
    assert f"> {AUDIT_NOTE_HUMAN_APPROVED}" in text

    # 无 audit_note 时不出现在输出。
    without = render_markdown(_sample_pack()).decode("utf-8")
    assert AUDIT_NOTE_HUMAN_APPROVED not in without


def test_render_docx_deterministic_and_reopenable() -> None:
    pack = _sample_pack(audit_note=AUDIT_NOTE_HUMAN_APPROVED)
    first = render_docx(pack)
    assert render_docx(pack) == first, "同 pack → 同 bytes（字节确定性）"
    assert first[:2] == b"PK", "docx 是 ZIP 容器"

    # reopen-able（python-docx）且中文正文可读。
    document = Document(BytesIO(first))
    texts = [p.text for p in document.paragraphs]
    assert "贵州茅台 基本面研究报告（600519）" in texts
    assert any("公司经营现金流净额同比增长20%。[1][2]" in t for t in texts)
    assert "证据附录" in texts
    assert any(AUDIT_NOTE_HUMAN_APPROVED in t for t in texts)


def test_render_docx_zip_timestamp_boundary_determinism() -> None:
    """A1/A3：同 pack 跨 ZIP 秒级时间边界 → 字节必须精确一致。

    当前实现只固定 core_properties.created/modified，python-docx 保存时
    ZipInfo.date_time 取**保存时刻**（秒级）——连续两次立即 render 落在同一秒
    会“脆弱通过”；跨过秒级边界 → 字节漂移，sha256 不等。本测试强制跨边界，
    同时校验每个 ZIP entry 的 date_time 固定为合法 ZIP 最小值（1980-01-01）。
    """
    pack = _sample_pack(audit_note=AUDIT_NOTE_HUMAN_APPROVED)

    first = render_docx(pack)
    with zipfile.ZipFile(BytesIO(first)) as zf:
        infos = zf.infolist()
    assert len(infos) >= 4, "OOXML 容器至少应有核心 XML part"
    # 所有 entry 的 ZIP timestamp 必须固定，不允许随保存时刻漂移。
    assert all(info.date_time == _ZIP_FIXED_TIMESTAMP for info in infos)

    # 跨过可见的秒级时间边界再渲染同一 pack → 必须得到完全相同的字节。
    now = time.monotonic()
    time.sleep(1.05 - (now % 1.0))
    second = render_docx(pack)
    assert second == first, "同 pack 跨时间边界 → 字节必须精确一致"

    # 重写不改 entry content：python-docx 重开成功，中文标题/正文/附录保留。
    document = Document(BytesIO(second))
    texts = [p.text for p in document.paragraphs]
    assert "贵州茅台 基本面研究报告（600519）" in texts
    assert any("公司经营现金流净额同比增长20%。[1][2]" in t for t in texts)
    assert "证据附录" in texts
    assert any(AUDIT_NOTE_HUMAN_APPROVED in t for t in texts)


def test_render_docx_different_pack_different_bytes() -> None:
    """不同 pack（正文/引用不同）→ docx 字节必须不同（normalize 不破坏区分度）。"""
    pack_a = _sample_pack()
    pack_b = _sample_pack(audit_note=AUDIT_NOTE_HUMAN_APPROVED)
    assert render_docx(pack_b) != render_docx(pack_a)


def test_render_pdf_deterministic_and_chinese_readable() -> None:
    pack = _sample_pack()
    first = render_pdf(pack)
    assert render_pdf(pack) == first, "同 pack → 同 bytes（字节确定性）"
    assert first[:4] == b"%PDF", "PDF magic"

    # pdfplumber 提取真实中文（spec O：PDF 中文可读）。
    import pdfplumber

    with pdfplumber.open(BytesIO(first)) as pdf:
        extracted = "\n".join(page.extract_text() or "" for page in pdf.pages)
    assert "贵州茅台" in extracted
    assert "公司经营现金流净额同比增长20%。[1][2]" in extracted


def test_render_export_registry_rejects_unknown_format() -> None:
    from app.report_export.errors import ReportExportError

    with pytest.raises(ReportExportError):
        render_export(_sample_pack(), "html")


# ---------------------------------------------------------------- fingerprint (spec M)


def test_fingerprint_stable_and_sensitive() -> None:
    pack = _sample_pack()
    base = compute_export_input_fingerprint(**_fingerprint_args(pack))
    # 同输入 → 同指纹（replay 依据）。
    assert compute_export_input_fingerprint(**_fingerprint_args(pack)) == base
    # format 变化 → 新指纹。
    assert compute_export_input_fingerprint(**_fingerprint_args(pack, format="pdf")) != base
    assert compute_export_input_fingerprint(**_fingerprint_args(pack, format="docx")) != base
    # report / check / audit / renderer / pack 任一变化 → 新指纹。
    assert (
        compute_export_input_fingerprint(**_fingerprint_args(pack, report_fingerprint="g" * 64))
        != base
    )
    assert compute_export_input_fingerprint(**_fingerprint_args(pack, audit_id=uuid4())) != base
    assert compute_export_input_fingerprint(**_fingerprint_args(pack, renderer_version=2)) != base
    assert (
        compute_export_input_fingerprint(
            **_fingerprint_args(_sample_pack(audit_note=AUDIT_NOTE_HUMAN_APPROVED))
        )
        != base
    )
    # 同输入但不同 check 指纹 → 新指纹（check 是导出身份的一部分）。
    assert (
        compute_export_input_fingerprint(**_fingerprint_args(pack, check_fingerprint="d" * 64))
        != base
    )


# ---------------------------------------------------------------- store (spec L)


def test_store_content_addressed_and_reused(tmp_path) -> None:
    store = ExportArtifactStore(tmp_path / "exports")
    data = "hello export 中文".encode()
    first = store.put_bytes(data, "md")
    second = store.put_bytes(data, "md")
    assert first.content_sha256 == hashlib.sha256(data).hexdigest()
    assert first.byte_size == len(data)
    assert first.storage_key == second.storage_key
    assert first.newly_created is True
    assert second.newly_created is False  # 相同内容复用同一文件
    assert store.exists(first.storage_key)
    with store.open(first.storage_key) as handle:
        assert handle.read() == data
    # 相同内容不同扩展名 → 不同 key。
    pdf = store.put_bytes(data, "pdf")
    assert pdf.storage_key != first.storage_key
    # 内容寻址路径位于根目录内。
    assert store._resolve(first.storage_key).is_relative_to((tmp_path / "exports").resolve())


def test_store_rejects_empty_and_path_traversal(tmp_path) -> None:
    store = ExportArtifactStore(tmp_path / "exports")
    with pytest.raises(ExportArtifactNotFound):
        store.put_bytes(b"", "md")
    for bad in (
        "/etc/passwd",
        "\\evil",
        "..",
        "../x",
        "a/../b",
        "a//b",
        "a/./b",
        "..\\..\\etc\\passwd",
        # 跨平台契约（regression）：反斜杠在 Linux 是普通字符、Windows 是分隔符，
        # storage key 必须 OS-independent —— 含 `\` 的 key 一律确定性拒绝，
        # 不得依赖 host Path 语义（Windows 曾靠 is_relative_to 兜底而 Linux 放行）。
        "a\\..\\b",
        "dir\\sub",
        "a/b\\..\\c",
        "..\\evil",
    ):
        with pytest.raises(InvalidStorageKey):
            store._resolve(bad)
    with pytest.raises(ExportArtifactNotFound):
        store.open("no/such/file.md")
    with pytest.raises(InvalidStorageKey):
        store.put_bytes(b"x", "txt")  # 扩展名白名单外
